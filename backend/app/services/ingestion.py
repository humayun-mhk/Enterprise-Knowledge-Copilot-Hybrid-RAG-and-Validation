from __future__ import annotations

import hashlib
import html
import logging
import re
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Iterable

from ..domain import ChunkRecord

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".html", ".htm"}
PAGE_MARKER_PATTERN = re.compile(
    r"^\s*Page\s+(\d+)(?:\s+of\s+\d+)?(?:\s*[|:—-]\s*Section\s*:\s*(.+))?\s*$",
    flags=re.IGNORECASE,
)


class UnsupportedDocumentError(ValueError):
    pass


class DocumentParseError(ValueError):
    pass


@dataclass(slots=True)
class ParsedBlock:
    text: str
    page_number: int | None = 1
    section: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(slots=True)
class ParsedDocument:
    blocks: list[ParsedBlock]
    page_count: int
    metadata: dict[str, Any] = field(default_factory=dict)


def clean_text(value: str) -> str:
    """Normalize extraction artifacts without destroying paragraph boundaries."""

    value = value.replace("\x00", "").replace("\u00ad", "")
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"(?<=\w)-\n(?=\w)", "", value)
    lines = [re.sub(r"[\t \u00a0]+", " ", line).strip() for line in value.split("\n")]
    output: list[str] = []
    blank = False
    for line in lines:
        if not line:
            if output and not blank:
                output.append("")
            blank = True
        else:
            output.append(line)
            blank = False
    return "\n".join(output).strip()


def is_probable_heading(line: str) -> bool:
    text = line.strip().strip(":")
    if not text or len(text) > 120 or len(text.split()) > 14:
        return False
    return bool(
        re.match(r"^(?:\d+(?:\.\d+)*[.)]?\s+|[A-Z][A-Z\s/&-]{3,})", text)
        or (text.istitle() and not text.endswith((".", "?", "!")))
    )


def split_sections(text: str, *, page_number: int | None, initial_section: str | None = None) -> list[ParsedBlock]:
    lines = text.splitlines()
    section = initial_section
    buffer: list[str] = []
    blocks: list[ParsedBlock] = []

    def flush() -> None:
        body = clean_text("\n".join(buffer))
        if body:
            blocks.append(ParsedBlock(text=body, page_number=page_number, section=section))
        buffer.clear()

    for line in lines:
        stripped = line.strip()
        if is_probable_heading(stripped) and buffer:
            flush()
            section = stripped.strip(":")
        elif is_probable_heading(stripped) and not buffer:
            section = stripped.strip(":")
        else:
            buffer.append(line)
    flush()
    return blocks


def parse_pdf(path: Path) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency error is environment-specific
        raise DocumentParseError("PDF support requires the 'pypdf' package") from exc
    try:
        reader = PdfReader(str(path))
        blocks: list[ParsedBlock] = []
        for index, page in enumerate(reader.pages, start=1):
            text = clean_text(page.extract_text() or "")
            if text:
                blocks.extend(split_sections(text, page_number=index))
        metadata = {}
        if reader.metadata:
            metadata = {
                str(key).lstrip("/").lower(): str(value)
                for key, value in dict(reader.metadata).items()
                if value is not None
            }
        return ParsedDocument(blocks=blocks, page_count=len(reader.pages), metadata=metadata)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Could not parse PDF: {exc}") from exc


def _docx_has_page_break(paragraph: Any) -> bool:
    try:
        return bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))
    except Exception:
        return False


def parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise DocumentParseError("DOCX support requires the 'python-docx' package") from exc
    try:
        document = Document(str(path))
        blocks: list[ParsedBlock] = []
        page = 1
        section: str | None = None
        buffer: list[str] = []

        def flush() -> None:
            text = clean_text("\n".join(buffer))
            if text:
                blocks.append(ParsedBlock(text=text, page_number=page, section=section))
            buffer.clear()

        for paragraph in document.paragraphs:
            text = clean_text(paragraph.text)
            style_name = str(getattr(paragraph.style, "name", "") or "")
            page_marker = PAGE_MARKER_PATTERN.match(text) if text else None
            if page_marker:
                flush()
                page = int(page_marker.group(1))
                if page_marker.group(2):
                    section = page_marker.group(2).strip()
            elif text and (style_name.lower().startswith("heading") or is_probable_heading(text)):
                flush()
                section = text
            elif text:
                buffer.append(text)
            if _docx_has_page_break(paragraph):
                flush()
                page += 1
        flush()
        for table_index, table in enumerate(document.tables, start=1):
            rows = [" | ".join(clean_text(cell.text) for cell in row.cells) for row in table.rows]
            table_text = clean_text("\n".join(rows))
            if table_text:
                blocks.append(
                    ParsedBlock(
                        text=table_text,
                        page_number=page,
                        section=section or f"Table {table_index}",
                        metadata={"content_type": "table", "table_index": table_index},
                    )
                )
        core = document.core_properties
        metadata = {
            key: value
            for key, value in {
                "title": core.title,
                "subject": core.subject,
                "author": core.author,
                "category": core.category,
                "keywords": core.keywords,
            }.items()
            if value
        }
        return ParsedDocument(blocks=blocks, page_count=page, metadata=metadata)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Could not parse DOCX: {exc}") from exc


class _StructuredHTMLParser(HTMLParser):
    block_tags = {"p", "li", "td", "th", "blockquote", "pre", "div", "section", "article"}
    heading_tags = {"h1", "h2", "h3", "h4", "h5", "h6", "title"}
    ignored_tags = {"script", "style", "noscript", "svg"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[tuple[int, str | None, str]] = []
        self.section: str | None = None
        self.page_number = 1
        self._tag: str | None = None
        self._buffer: list[str] = []
        self._ignored_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in self.ignored_tags:
            self._ignored_depth += 1
            return
        if self._ignored_depth:
            return
        attributes = dict(attrs)
        if tag == "article" and attributes.get("data-page-number"):
            try:
                self.page_number = int(str(attributes["data-page-number"]))
            except ValueError:
                pass
            if attributes.get("data-section"):
                self.section = str(attributes["data-section"])
        if tag in self.block_tags | self.heading_tags:
            self._flush()
            self._tag = tag
        elif tag == "br":
            self._buffer.append("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.ignored_tags:
            self._ignored_depth = max(0, self._ignored_depth - 1)
            return
        if not self._ignored_depth and tag in self.block_tags | self.heading_tags:
            self._flush()

    def handle_data(self, data: str) -> None:
        if not self._ignored_depth:
            self._buffer.append(data)

    def _flush(self) -> None:
        text = clean_text(html.unescape(" ".join(self._buffer)))
        self._buffer.clear()
        if not text:
            self._tag = None
            return
        if self._tag in self.heading_tags:
            self.section = text
        else:
            self.blocks.append((self.page_number, self.section, text))
        self._tag = None

    def close(self) -> None:
        super().close()
        self._flush()


def parse_html(path: Path) -> ParsedDocument:
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
        parser = _StructuredHTMLParser()
        parser.feed(raw)
        parser.close()
        blocks = [
            ParsedBlock(text=text, page_number=page_number, section=section)
            for page_number, section, text in parser.blocks
        ]
        if not blocks:
            stripped = clean_text(re.sub(r"<[^>]+>", " ", raw))
            if stripped:
                blocks = [ParsedBlock(text=stripped, page_number=1)]
        return ParsedDocument(
            blocks=blocks,
            page_count=max((block.page_number or 1 for block in blocks), default=1),
            metadata={},
        )
    except Exception as exc:
        raise DocumentParseError(f"Could not parse HTML: {exc}") from exc


def parse_txt(path: Path) -> ParsedDocument:
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
        pages = raw.split("\f")
        blocks: list[ParsedBlock] = []
        for index, page_text in enumerate(pages, start=1):
            text = clean_text(page_text)
            if text:
                blocks.extend(split_sections(text, page_number=index))
        return ParsedDocument(blocks=blocks, page_count=max(1, len(pages)), metadata={})
    except Exception as exc:
        raise DocumentParseError(f"Could not parse text document: {exc}") from exc


def parse_document(path: Path, extension: str | None = None) -> ParsedDocument:
    suffix = (extension or path.suffix).lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentError(
            f"Unsupported document type '{suffix}'. Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    parser = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".txt": parse_txt,
        ".html": parse_html,
        ".htm": parse_html,
    }[suffix]
    parsed = parser(path)
    if not parsed.blocks:
        raise DocumentParseError("The document did not contain extractable text")
    return parsed


def chunk_text(text: str, *, chunk_size: int, overlap: int) -> Iterable[str]:
    if chunk_size < 1:
        raise ValueError("chunk_size must be positive")
    overlap = max(0, min(overlap, chunk_size - 1))
    text = clean_text(text)
    start = 0
    length = len(text)
    while start < length:
        target_end = min(length, start + chunk_size)
        end = target_end
        if target_end < length:
            floor = start + max(chunk_size // 2, 1)
            candidates = [text.rfind("\n", floor, target_end), text.rfind(". ", floor, target_end), text.rfind(" ", floor, target_end)]
            boundary = max(candidates)
            if boundary > start:
                end = boundary + (1 if text[boundary : boundary + 2] == ". " else 0)
        piece = text[start:end].strip()
        if piece:
            yield piece
        if end >= length:
            break
        next_start = max(start + 1, end - overlap)
        while next_start < length and next_start > 0 and text[next_start - 1].isalnum() and text[next_start].isalnum():
            next_start += 1
        start = next_start


def build_chunks(
    parsed: ParsedDocument,
    *,
    document_id: str,
    document_name: str,
    chunk_size: int,
    overlap: int,
) -> tuple[list[ChunkRecord], int]:
    chunks: list[ChunkRecord] = []
    seen_hashes: set[str] = set()
    duplicates = 0
    sequence = 0
    for block in parsed.blocks:
        for text in chunk_text(block.text, chunk_size=chunk_size, overlap=overlap):
            content_hash = hashlib.sha256(text.casefold().encode("utf-8")).hexdigest()
            if content_hash in seen_hashes:
                duplicates += 1
                continue
            seen_hashes.add(content_hash)
            sequence += 1
            stable = f"{document_id}:{block.page_number}:{sequence}:{content_hash}".encode("utf-8")
            chunk_id = "chunk_" + hashlib.sha256(stable).hexdigest()[:20]
            chunks.append(
                ChunkRecord(
                    chunk_id=chunk_id,
                    document_id=document_id,
                    document_name=document_name,
                    page_number=block.page_number,
                    section=block.section,
                    chunk_text=text,
                    metadata={
                        **block.metadata,
                        "sequence": sequence,
                        "document_metadata": parsed.metadata,
                    },
                    content_hash=content_hash,
                )
            )
    return chunks, duplicates


def file_checksum(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def safe_filename(filename: str) -> str:
    name = Path(filename or "document").name
    stem = re.sub(r"[^A-Za-z0-9._ -]+", "_", name).strip(" .")
    return stem[:180] or "document.txt"
